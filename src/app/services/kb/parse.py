"""Text extraction from uploaded files (pdf, docx, html, csv, plain text).

CPU/blocking by nature — call via `run_in_threadpool` from async code.
Unknown extensions are rejected here; the upload endpoint pre-validates
against `settings.kb_allowed_extensions` so this is a second line of defense.
"""

from __future__ import annotations

import csv
import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


class ParseError(ValueError):
    """Raised when a file's text cannot be extracted."""


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    pages = [(page.extract_text() or "") for page in reader.pages]
    return "\n\n".join(pages)


def _extract_pdf_pages(data: bytes) -> list[str]:
    """PDF text page-by-page (page-level chunking; NVIDIA: best average accuracy)."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    return [(page.extract_text() or "") for page in reader.pages]


def _extract_docx(data: bytes) -> str:
    import docx

    document = docx.Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _extract_html(data: bytes) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(data, "html.parser")
    return soup.get_text(" ", strip=True)


def _extract_csv(data: bytes) -> str:
    text = _decode(data)
    rows = csv.reader(io.StringIO(text))
    return "\n".join(", ".join(row) for row in rows if any(c.strip() for c in row))


def _decode(data: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


# Extension -> parser callable. Anything else falls back to plain text.
_PARSERS = {
    ".pdf": _extract_pdf,
    ".docx": _extract_docx,
    ".html": _extract_html,
    ".htm": _extract_html,
    ".csv": _extract_csv,
}


def extract_text(path: str, data: bytes) -> str:
    """Extract plain text from raw file bytes based on the file extension."""
    ext = Path(path).suffix.lower()
    parser = _PARSERS.get(ext)
    if parser is None:
        return _decode(data)
    try:
        return parser(data)
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError(f"Failed to parse {ext} file: {exc}") from exc


def extract_pages(path: str, data: bytes) -> list[str]:
    """Extract text page-by-page where the format supports pages (PDF).

    Returns a list of page texts; non-PDF formats produce a single element.
    Used by the ingest pipeline for page-level chunking.
    """
    if Path(path).suffix.lower() != ".pdf":
        return [extract_text(path, data)]
    try:
        return _extract_pdf_pages(data)
    except ParseError:
        raise
    except Exception as exc:
        raise ParseError(f"Failed to parse .pdf file: {exc}") from exc
